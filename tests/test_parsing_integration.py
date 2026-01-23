import json
import os
import unittest

import docx

from app import app


class TestParsingIntegration(unittest.TestCase):
    def setUp(self):
        self.app = app.test_client()
        self.app.testing = True

        # Create a dummy user session (username as string)
        with self.app.session_transaction() as sess:
            sess["user"] = "testuser"

        # Ensure user directory exists
        self.user_dir = os.path.join("data", "testuser")
        os.makedirs(self.user_dir, exist_ok=True)

        # Create dummy docx
        self.docx_path = "test_resume.docx"
        doc = docx.Document()
        doc.add_heading("John Doe", 0)
        doc.add_heading("Experience", level=1)
        doc.add_paragraph("Software Engineer at Tech Corp")
        doc.save(self.docx_path)

    def tearDown(self):
        if os.path.exists(self.docx_path):
            os.remove(self.docx_path)
        # Clean up user dir if needed, but risky if real data exists.
        # Since it's 'testuser', it should be safe to leave or clean.
        # I'll leave it to avoid accidental deletion of important things.

    def test_upload_docx(self):
        with open(self.docx_path, "rb") as f:
            data = {"file": (f, "test_resume.docx")}
            response = self.app.post("/api/upload_resume", data=data, content_type="multipart/form-data")

        self.assertEqual(response.status_code, 200)
        data = json.loads(response.data)
        self.assertEqual(data["status"], "success")
        self.assertIn("John Doe", data["text"])
        self.assertIn("Software Engineer", data["text"])


if __name__ == "__main__":
    unittest.main()
